import html
import io
import os
import re
import sys
import unicodedata
from datetime import datetime
from html.parser import HTMLParser

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from markdown.extensions.fenced_code import FencedCodeExtension
from markdown.extensions.tables import TableExtension
import markdown
from pydantic import BaseModel
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from xhtml2pdf import pisa

app = FastAPI(
    title="DocMind Core API",
    description="API de Conversão Direta de Markdown para PDF e DOCX - Mindworks",
    version="2.0.0",
)

COR_NAVY = "003366"
COR_ROYAL = "1D4ED8"
COR_CINZA_CLARO = "666666"

RGB_NAVY = RGBColor(0x00, 0x33, 0x66)
RGB_ROYAL = RGBColor(0x1D, 0x4E, 0xD8)
RGB_CINZA_CLARO = RGBColor(0x66, 0x66, 0x66)

FONTE_PADRAO = "Arial"
FONTE_CODIGO = "Courier New"

TITULOS_POR_TIPO = {
    "arquitetura-backend": "Documentação Técnica - Backend",
    "estrutura-frontend": "Documentação Técnica - Frontend",
    "visao-geral-cliente": "Visão Geral do Projeto",
    "backend": "Documentação Técnica - Backend",
    "frontend": "Documentação Técnica - Frontend",
    "geral": "Visão Geral do Projeto",
}


class DocumentPayload(BaseModel):
    texto: str
    tipoDocumentacao: str = "visao-geral-cliente"
    tipo: str | None = None
    titulo: str | None = None


def resolver_titulo(payload: DocumentPayload) -> str:
    if payload.titulo:
        return payload.titulo
    tipo_escolhido = payload.tipo if payload.tipo else payload.tipoDocumentacao
    return TITULOS_POR_TIPO.get(tipo_escolhido, "Visão Geral do Projeto")


def sanitizar_nome_arquivo(nome: str) -> str:
    nfkd = unicodedata.normalize("NFKD", nome)
    palavra_limpa = "".join([c for c in nfkd if not unicodedata.combining(c)])
    return re.sub(r"[^a-zA-Z0-9-_]", "_", palavra_limpa).lower()


def remover_titulo_duplicado(texto_markdown: str) -> str:
    linhas = texto_markdown.lstrip("\n").split("\n")
    if linhas and linhas[0].startswith("# "):
        linhas = linhas[1:]
        if linhas and linhas[0].strip() == "":
            linhas = linhas[1:]
    return "\n".join(linhas)


def sanitizar_html_bruto(texto_markdown: str) -> str:
    return html.escape(texto_markdown, quote=False)


_FONTE_REGISTRADA = None


def _registrar_fonte_arial() -> str:
    global _FONTE_REGISTRADA
    if _FONTE_REGISTRADA:
        return _FONTE_REGISTRADA

    diretorio_script = os.path.dirname(os.path.abspath(__file__))

    candidatos = [
        (
            os.path.join(diretorio_script, "fonts", "arial.ttf"),
            os.path.join(diretorio_script, "fonts", "arialbd.ttf"),
        ),
        (
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        ),
        (
            "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
            "/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf",
        ),
        (
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\arialbd.ttf",
        ),
    ]

    for caminho_regular, caminho_bold in candidatos:
        try:
            if os.path.isfile(caminho_regular) and os.path.isfile(caminho_bold):
                pdfmetrics.registerFont(TTFont("Arial", caminho_regular))
                pdfmetrics.registerFont(TTFont("Arial-Bold", caminho_bold))
                _FONTE_REGISTRADA = "Arial"
                return _FONTE_REGISTRADA
        except Exception:
            continue

    print(
        "[AVISO] Nenhum arquivo de fonte Arial (ou equivalente) foi encontrado nos "
        "caminhos verificados. O PDF será gerado com Helvetica, o que pode ficar "
        "visualmente diferente do DOCX (que sempre referencia 'Arial'). Para "
        "garantir consistência, adicione arial.ttf e arialbd.ttf na pasta "
        "'fonts/' ao lado deste script.",
        file=sys.stderr,
    )
    _FONTE_REGISTRADA = "Helvetica"
    return _FONTE_REGISTRADA


def _construir_css(fonte: str) -> str:
    return f"""
@page {{
    size: a4 portrait;
    margin: 2.5cm 2cm 2.5cm 2cm;
    @frame footer_frame {{
        -pdf-frame-content: footer_content;
        bottom: 1cm;
        margin-left: 2cm;
        margin-right: 2cm;
        height: 1cm;
    }}
}}
body {{
    font-family: {fonte}, sans-serif;
    font-size: 10pt;
    line-height: 1.45;
    color: #222222;
}}
#cabecalho {{
    border-bottom: 2pt solid #{COR_NAVY};
    padding-bottom: 12px;
    margin-bottom: 20px;
}}
#cabecalho .marca {{
    color: #{COR_ROYAL};
    font-size: 8.5pt;
    font-weight: bold;
    margin: 0 0 4px 0;
}}
#cabecalho h1 {{
    color: #{COR_NAVY};
    font-size: 16pt;
    margin: 0;
}}
#cabecalho .data {{
    color: #{COR_CINZA_CLARO};
    font-size: 8pt;
    margin: 4px 0 0 0;
}}
h1, h2, h3 {{ color: #{COR_NAVY}; page-break-after: avoid; }}
h1 {{ font-size: 14pt; border-bottom: 1pt solid #{COR_NAVY}; margin-top: 18px; }}
h2 {{ font-size: 12pt; margin-top: 14px; }}
h3 {{ font-size: 10pt; margin-top: 10px; }}
p {{ margin: 6px 0; text-align: justify; }}
code {{ font-family: Courier, monospace; background-color: #f0f2f5; padding: 1px 4px; }}
pre {{ background-color: #f5f6f8; border-left: 3px solid #{COR_NAVY}; padding: 8px 12px; font-family: Courier, monospace; white-space: pre-wrap; }}
table {{
    border-collapse: collapse;
    width: 100%;
    margin: 12px 0;
    font-size: 9pt;
    table-layout: fixed;
}}
th, td {{
    border: 1pt solid #000000;
    padding: 6px 8px;
    text-align: left;
    word-wrap: break-word;
    overflow: hidden;
}}
tr {{ page-break-inside: avoid; }}
th {{ background-color: #{COR_NAVY}; color: #ffffff; }}
#footer_content {{ font-size: 8pt; color: #{COR_CINZA_CLARO}; text-align: center; border-top: 0.5pt solid #d0d5dd; padding-top: 4px; }}
"""


def _inserir_numeracao_html(html_content: str) -> str:
    niveis_presentes = sorted({int(nivel) for nivel in re.findall(r"<h([123])>", html_content)})

    if not niveis_presentes:
        return html_content

    nivel_principal = niveis_presentes[0]
    contadores = {1: 0, 2: 0, 3: 0}

    def replace_match(match):
        tag = match.group(1)
        content = match.group(2)
        nivel_atual = int(tag[1])

        for nivel in range(nivel_atual + 1, 4):
            contadores[nivel] = 0

        contadores[nivel_atual] += 1

        partes_numeracao = [
            str(contadores[nivel] or 1)
            for nivel in range(nivel_principal, nivel_atual + 1)
        ]
        prefixo = ".".join(partes_numeracao) + ". "

        return f"<{tag}>{prefixo}{content}</{tag}>"

    return re.sub(r'<(h[123])>(.*?)</\1>', replace_match, html_content)


def _corrigir_indentacao_sublistas(texto_markdown: str) -> str:
    linhas = texto_markdown.split("\n")
    resultado = []
    dentro_de_sublista = False
    padrao_label_com_dois_pontos = re.compile(r'^-\s+\*\*[^*]+\*\*:\s*$')
    padrao_item_lista = re.compile(r'^-\s+\S')

    for linha in linhas:
        linha_stripped = linha.strip()

        if padrao_label_com_dois_pontos.match(linha_stripped):
            resultado.append(linha_stripped)
            dentro_de_sublista = True
            continue

        if dentro_de_sublista:
            if linha_stripped == "":
                dentro_de_sublista = False
                resultado.append(linha)
                continue
            if padrao_item_lista.match(linha_stripped):
                resultado.append("    " + linha_stripped)
                continue
            dentro_de_sublista = False
            resultado.append(linha)
            continue

        resultado.append(linha)

    return "\n".join(resultado)


def _markdown_para_html(texto_markdown: str) -> str:
    texto_sanitizado = sanitizar_html_bruto(texto_markdown)
    texto_sem_titulo = remover_titulo_duplicado(texto_sanitizado)
    texto_corrigido = re.sub(r'(?m)^(\s*\d+\.)\s*\n+\s*', r'\1 ', texto_sem_titulo)
    texto_com_sublistas_corrigidas = _corrigir_indentacao_sublistas(texto_corrigido)
    html_conteudo = markdown.markdown(
        texto_com_sublistas_corrigidas, extensions=[TableExtension(), FencedCodeExtension()]
    )
    return _inserir_numeracao_html(html_conteudo)


def _aplicar_fonte_run(run, nome_fonte: str = FONTE_PADRAO):
    run.font.name = nome_fonte
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), nome_fonte)
    r_fonts.set(qn("w:hAnsi"), nome_fonte)
    r_fonts.set(qn("w:eastAsia"), nome_fonte)
    r_fonts.set(qn("w:cs"), nome_fonte)


def _aplicar_fonte_estilo(doc: Document, nome_estilo: str, nome_fonte: str = FONTE_PADRAO):
    try:
        estilo = doc.styles[nome_estilo]
    except KeyError:
        return

    estilo.font.name = nome_fonte
    try:
        r_pr = estilo.element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), nome_fonte)
        r_fonts.set(qn("w:hAnsi"), nome_fonte)
        r_fonts.set(qn("w:eastAsia"), nome_fonte)
        r_fonts.set(qn("w:cs"), nome_fonte)
    except Exception:
        pass


def _padronizar_fontes_do_documento(doc: Document):
    for nome_estilo in (
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Bullet",
        "List Bullet 2",
        "List Bullet 3",
        "List Number",
        "List Number 2",
        "List Number 3",
    ):
        _aplicar_fonte_estilo(doc, nome_estilo, FONTE_PADRAO)


def _inserir_campo_word(paragrafo, instrucao_campo: str):
    run = paragrafo.add_run()
    _aplicar_fonte_run(run, FONTE_PADRAO)
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instrucao = OxmlElement("w:instrText")
    instrucao.set(qn("xml:space"), "preserve")
    instrucao.text = instrucao_campo
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    fim = OxmlElement("w:fldChar")
    fim.set(qn("w:fldCharType"), "end")
    run._r.append(inicio)
    run._r.append(instrucao)
    run._r.append(separador)
    run._r.append(fim)
    return run


def adicionar_cabecalho_rodape_docx(doc: Document, titulo_documento: str):
    marca = doc.add_paragraph()
    r_marca = marca.add_run("MINDWORKS")
    _aplicar_fonte_run(r_marca, FONTE_PADRAO)
    r_marca.font.size = Pt(8.5)
    r_marca.font.bold = True
    r_marca.font.color.rgb = RGB_ROYAL

    titulo = doc.add_paragraph()
    r_titulo = titulo.add_run(titulo_documento)
    _aplicar_fonte_run(r_titulo, FONTE_PADRAO)
    r_titulo.font.size = Pt(16)
    r_titulo.font.bold = True
    r_titulo.font.color.rgb = RGB_NAVY

    data = doc.add_paragraph()
    r_data = data.add_run(f"Data de Emissão: {datetime.now().strftime('%d/%m/%Y')}")
    _aplicar_fonte_run(r_data, FONTE_PADRAO)
    r_data.font.size = Pt(8.5)
    r_data.font.color.rgb = RGB_CINZA_CLARO

    paragrafo = doc.sections[0].footer.paragraphs[0]
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_footer = paragrafo.add_run("Página ")
    _aplicar_fonte_run(r_footer, FONTE_PADRAO)
    r_footer.font.size = Pt(8)

    _inserir_campo_word(paragrafo, "PAGE")

    r_footer2 = paragrafo.add_run(" de ")
    _aplicar_fonte_run(r_footer2, FONTE_PADRAO)
    r_footer2.font.size = Pt(8)

    _inserir_campo_word(paragrafo, "NUMPAGES")

    r_footer3 = paragrafo.add_run("    ·    Mindworks - Documento Confidencial")
    _aplicar_fonte_run(r_footer3, FONTE_PADRAO)
    r_footer3.font.size = Pt(8)


_PROXIMO_NUM_ID = [9000]


def _garantir_abstract_num_lista_numerada(doc: Document) -> str:
    numbering_element = doc.part.numbering_part.element
    abstract_num_id_fixo = "9000"

    for abstract_existente in numbering_element.findall(qn("w:abstractNum")):
        if abstract_existente.get(qn("w:abstractNumId")) == abstract_num_id_fixo:
            return abstract_num_id_fixo

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_num_id_fixo)

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    abstract_num.append(lvl)

    numbering_element.insert(0, abstract_num)
    return abstract_num_id_fixo


def _criar_nova_lista_numerada(doc: Document) -> str:
    abstract_num_id = _garantir_abstract_num_lista_numerada(doc)
    numbering_element = doc.part.numbering_part.element

    novo_num_id = str(_PROXIMO_NUM_ID[0])
    _PROXIMO_NUM_ID[0] += 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), novo_num_id)

    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_num_id)
    num.append(abstract_ref)

    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)

    numbering_element.append(num)

    return novo_num_id


def _aplicar_numeracao_paragrafo(paragrafo, num_id: str):
    p_pr = paragrafo._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def _obter_estilo_lista_por_profundidade(nome_base: str, profundidade: int) -> str:
    profundidade_limitada = min(max(profundidade, 1), 3)
    if profundidade_limitada == 1:
        return nome_base
    return f"{nome_base} {profundidade_limitada}"


class ConversorHtmlParaDocx(HTMLParser):

    TAGS_BLOCO_IGNORADAS = {"html", "body", "div"}

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc

        self.paragrafo_atual = None
        self.pilha_formatacao = []

        self.pilha_listas = []
        self.pilha_num_ids = []
        self.num_id_lista_atual = None

        self.li_recem_criado = False

        self.em_tabela = False
        self.linhas_tabela = []
        self.linha_atual = []
        self.celula_atual_textos = []
        self.em_cabecalho_tabela = False

        self.em_pre = False
        self.buffer_pre = []

    def _iniciar_heading(self, nivel: int):
        estilo = f"Heading {nivel}"
        self.paragrafo_atual = self.doc.add_paragraph(style=estilo)

    def _finalizar_heading_cor(self):
        for run in self.paragrafo_atual.runs:
            _aplicar_fonte_run(run, FONTE_PADRAO)
            run.font.color.rgb = RGB_NAVY
            run.font.bold = True

    def handle_starttag(self, tag, attrs):
        if tag in ("h1", "h2", "h3", "h4"):
            nivel = int(tag[1])
            self._iniciar_heading(min(nivel, 3))
            self.pilha_listas = []
            self.pilha_num_ids = []
            self.num_id_lista_atual = None
        elif tag == "p":
            if self.li_recem_criado:
                self.li_recem_criado = False
            else:
                self.paragrafo_atual = self.doc.add_paragraph()
        elif tag in ("strong", "b"):
            self.pilha_formatacao.append("strong")
        elif tag in ("em", "i"):
            self.pilha_formatacao.append("em")
        elif tag == "code" and not self.em_pre:
            self.pilha_formatacao.append("code")
        elif tag == "pre":
            self.em_pre = True
            self.buffer_pre = []
        elif tag == "ul":
            self.pilha_listas.append("ul")
            self.pilha_num_ids.append(None)
        elif tag == "ol":
            self.pilha_listas.append("ol")
            novo_num_id = _criar_nova_lista_numerada(self.doc)
            self.pilha_num_ids.append(novo_num_id)
            self.num_id_lista_atual = novo_num_id
        elif tag == "li":
            self.li_recem_criado = True
            lista_atual = self.pilha_listas[-1] if self.pilha_listas else None
           
            profundidade = len(self.pilha_listas)
            if lista_atual == "ul":
                estilo_bullet = _obter_estilo_lista_por_profundidade("List Bullet", profundidade)
                self.paragrafo_atual = self.doc.add_paragraph(style=estilo_bullet)
            elif lista_atual == "ol":
                estilo_numero = _obter_estilo_lista_por_profundidade("List Number", profundidade)
                self.paragrafo_atual = self.doc.add_paragraph(style=estilo_numero)
                _aplicar_numeracao_paragrafo(self.paragrafo_atual, self.num_id_lista_atual)
            else:
                self.paragrafo_atual = self.doc.add_paragraph()
        elif tag == "table":
            self.em_tabela = True
            self.linhas_tabela = []
        elif tag == "thead":
            self.em_cabecalho_tabela = True
        elif tag == "tbody":
            self.em_cabecalho_tabela = False
        elif tag == "tr":
            self.linha_atual = []
        elif tag in ("td", "th"):
            self.celula_atual_textos = []
        elif tag == "br":
            if self.paragrafo_atual is not None:
                self.paragrafo_atual.add_run().add_break()

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "h4"):
            self._finalizar_heading_cor()
            self.paragrafo_atual = None
        elif tag == "p":
            self.paragrafo_atual = None
        elif tag in ("strong", "b"):
            if "strong" in self.pilha_formatacao:
                self.pilha_formatacao.remove("strong")
        elif tag in ("em", "i"):
            if "em" in self.pilha_formatacao:
                self.pilha_formatacao.remove("em")
        elif tag == "code" and not self.em_pre:
            if "code" in self.pilha_formatacao:
                self.pilha_formatacao.remove("code")
        elif tag == "pre":
            self.em_pre = False
            paragrafo = self.doc.add_paragraph()
            run = paragrafo.add_run("".join(self.buffer_pre))
            _aplicar_fonte_run(run, FONTE_CODIGO)
            run.font.size = Pt(9)
            self.buffer_pre = []
        elif tag in ("ul", "ol"):
            if self.pilha_listas:
                self.pilha_listas.pop()
            if self.pilha_num_ids:
                self.pilha_num_ids.pop()
            self.num_id_lista_atual = self.pilha_num_ids[-1] if self.pilha_num_ids else None
        elif tag == "li":
            self.paragrafo_atual = None
            self.li_recem_criado = False
        elif tag in ("td", "th"):
            self.linha_atual.append(("".join(self.celula_atual_textos), tag == "th" or self.em_cabecalho_tabela))
            self.celula_atual_textos = []
        elif tag == "tr":
            if self.linha_atual:
                self.linhas_tabela.append(self.linha_atual)
            self.linha_atual = []
        elif tag == "table":
            self._renderizar_tabela()
            self.em_tabela = False
            self.linhas_tabela = []

    def handle_data(self, data):
        if self.em_pre:
            self.buffer_pre.append(data)
            return

        if self.em_tabela:
            self.celula_atual_textos.append(data)
            return

        if not data.strip():
            return

        if self.paragrafo_atual is None:
            self.paragrafo_atual = self.doc.add_paragraph()

        run = self.paragrafo_atual.add_run(data)
        if "code" in self.pilha_formatacao:
            _aplicar_fonte_run(run, FONTE_CODIGO)
        else:
            _aplicar_fonte_run(run, FONTE_PADRAO)
        run.font.bold = "strong" in self.pilha_formatacao
        run.font.italic = "em" in self.pilha_formatacao
        if not self.paragrafo_atual.style.name.startswith("Heading"):
            run.font.size = Pt(10)

    def _renderizar_tabela(self):
        if not self.linhas_tabela:
            return

        num_linhas = len(self.linhas_tabela)
        num_colunas = max(len(linha) for linha in self.linhas_tabela)

        tabela = self.doc.add_table(rows=num_linhas, cols=num_colunas)
        tabela.style = "Table Grid"

        for i, linha in enumerate(self.linhas_tabela):
            for j, (texto_celula, eh_cabecalho) in enumerate(linha):
                celula = tabela.cell(i, j)
                celula.text = ""
                _definir_borda_celula(celula)
                paragrafo = celula.paragraphs[0]
                run = paragrafo.add_run(texto_celula.strip())
                _aplicar_fonte_run(run, FONTE_PADRAO)
                run.font.size = Pt(9)
                if eh_cabecalho:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _definir_cor_fundo_celula(celula, COR_NAVY)

        self.doc.add_paragraph()


def _definir_cor_fundo_celula(celula, cor_hex: str):
    propriedades_celula = celula._tc.get_or_add_tcPr()
    sombreamento = OxmlElement("w:shd")
    sombreamento.set(qn("w:fill"), cor_hex)
    propriedades_celula.append(sombreamento)


COR_BORDA_TABELA = "000000"


def _definir_borda_celula(celula, cor_hex: str = COR_BORDA_TABELA, tamanho: int = 4):
    propriedades_celula = celula._tc.get_or_add_tcPr()
    bordas = OxmlElement("w:tcBorders")
    for nome_borda in ("top", "left", "bottom", "right"):
        elemento_borda = OxmlElement(f"w:{nome_borda}")
        elemento_borda.set(qn("w:val"), "single")
        elemento_borda.set(qn("w:sz"), str(tamanho))
        elemento_borda.set(qn("w:space"), "0")
        elemento_borda.set(qn("w:color"), cor_hex)
        bordas.append(elemento_borda)
    propriedades_celula.append(bordas)


@app.post("/api/gerar-pdf")
async def gerar_pdf(payload: DocumentPayload):
    fonte = _registrar_fonte_arial()

    html_conteudo = _markdown_para_html(payload.texto)

    titulo_documento = html.escape(resolver_titulo(payload))
    data_geracao = f"Data de Emissão: {datetime.now().strftime('%d/%m/%Y')}"

    html_completo = f"""
        <html>
        <head><meta charset="utf-8"><style>{_construir_css(fonte)}</style></head>
        <body>
            <div id="cabecalho">
                <p class="marca">MINDWORKS</p>
                <h1>{titulo_documento}</h1>
                <p class="data">{data_geracao}</p>
            </div>
            <div id="footer_content">
                Página <pdf:pagenumber/> de <pdf:pagecount/> &bull; Mindworks - Documento Confidencial
            </div>
            {html_conteudo}
        </body>
        </html>
    """

    pdf_buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(html_completo, dest=pdf_buffer)
    if pisa_status.err:
        raise HTTPException(status_code=500, detail="Erro ao gerar o arquivo PDF.")

    pdf_buffer.seek(0)
    nome_arquivo = sanitizar_nome_arquivo(resolver_titulo(payload))

    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{nome_arquivo}.pdf"'
        },
    )


@app.post("/api/gerar-doc")
async def gerar_doc(payload: DocumentPayload):
    titulo_documento = resolver_titulo(payload)
    doc = Document()

    _padronizar_fontes_do_documento(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    adicionar_cabecalho_rodape_docx(doc, titulo_documento)

    html_conteudo = _markdown_para_html(payload.texto)

    conversor = ConversorHtmlParaDocx(doc)
    conversor.feed(html_conteudo)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nome_arquivo = sanitizar_nome_arquivo(titulo_documento)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{nome_arquivo}.docx"'
        },
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)