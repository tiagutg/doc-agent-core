import os
import json
import io
import zipfile
import markdown
from fastapi import FastAPI, HTTPException, UploadFile, File, Header
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List
from xhtml2pdf import pisa
from docx import Document
from markdown.extensions.tables import TableExtension
from markdown.extensions.fenced_code import FencedCodeExtension
from google.genai import types

# Importações da sua lógica de IA
from src.agents.doc_agent import criar_agente_analisador_arquivos, estruturar_tarefa_analise
from src.config.schema import AnaliseArquivoJSON
from src.agents.consolidator_agent import criar_agente_consolidador, estruturar_tarefa_consolidacao
from src.agents.tech_doc_agent import criar_agente_documentador_tecnico, estruturar_tarefa_documento_final

app = FastAPI(
    title="Doc Agent Core API",
    description="API robusta para documentação automatizada (Suporte a arquivos individuais e ZIP)",
    version="2.0.0"
)

# --- MODELOS ---
class RequisicaoAnalise(BaseModel):
    nome_arquivo: str
    conteudo_codigo: str

class RequisicaoConsolidacao(BaseModel):
    relatorios_json: List[str]

# --- ROTA: ANÁLISE DE PROJETO (ZIP) ---
@app.post("/analisar-zip")
async def processar_zip(file: UploadFile = File(...)):
    try:
        buffer = io.BytesIO(await file.read())
        resultados = []
        with zipfile.ZipFile(buffer) as z:
            # Lista arquivos, ignorando pastas de sistema e node_modules
            for nome_arquivo in z.namelist():
                if not nome_arquivo.endswith('/') and not any(x in nome_arquivo for x in ['node_modules', '.git', '__pycache__']):
                    with z.open(nome_arquivo) as f:
                        conteudo = f.read().decode("utf-8", errors="ignore")
                        client, sistema = criar_agente_analisador_arquivos()
                        resposta = client.models.generate_content(
                            model='gemini-1.5-pro',
                            contents=estruturar_tarefa_analise(nome_arquivo, conteudo),
                            config=types.GenerateContentConfig(
                                system_instruction=sistema,
                                response_mime_type="application/json",
                                response_schema=AnaliseArquivoJSON
                            ),
                        )
                        resultados.append(json.loads(resposta.text))
        return {"status": "sucesso", "analises": resultados}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro no processamento ZIP: {str(e)}")

# --- ROTA: FLUXO ANTIGO (ARQUIVO ÚNICO) ---
@app.post("/api/analisar")
async def analisar_arquivo(dados: RequisicaoAnalise):
    try:
        client, sistema = criar_agente_analisador_arquivos()
        resposta = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=estruturar_tarefa_analise(dados.nome_arquivo, dados.conteudo_codigo),
            config=types.GenerateContentConfig(
                system_instruction=sistema,
                response_mime_type="application/json",
                response_schema=AnaliseArquivoJSON
            ),
        )
        return json.loads(resposta.text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- ROTA: CONSOLIDAÇÃO ---
@app.post("/api/documentar")
async def consolidar_e_documentar(dados: RequisicaoConsolidacao):
    try:
        dados_base = "\n\n".join([f"--- COMPONENTE ---\n{r}" for r in dados.relatorios_json])
        
        # Consolidação
        client_c, sistema_c = criar_agente_consolidador()
        resposta_c = client_c.models.generate_content(model='gemini-2.5-flash', contents=estruturar_tarefa_consolidacao(dados_base), config={'system_instruction': sistema_c})
        
        # Documentação Final
        client_t, sistema_t = criar_agente_documentador_tecnico()
        resposta_t = client_t.models.generate_content(model='gemini-2.5-flash', contents=estruturar_tarefa_documento_final(resposta_c.text))
        
        return {"documentacao_final_markdown": resposta_t.text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- ROTAS DE EXPORTAÇÃO (PDF/DOCX) ---
@app.post("/api/gerar-pdf")
async def gerar_pdf(payload: dict):
    html_puro = markdown.markdown(payload.get("texto", ""), extensions=[TableExtension(), FencedCodeExtension()])
    pdf_buffer = io.BytesIO()
    pisa.CreatePDF(f"<html><body>{html_puro}</body></html>", dest=pdf_buffer)
    pdf_buffer.seek(0)
    return StreamingResponse(pdf_buffer, media_type="application/pdf")

@app.post("/api/gerar-doc")
async def gerar_doc(payload: dict):
    doc = Document()
    doc.add_paragraph(payload.get("texto", ""))
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)